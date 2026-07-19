"""Completa el DOCX convertido por Pandoc con portada, índice y formato APA 7."""

import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph


TITULO = (
    "Mundial Trivia Challenge: Sistema Interactivo de Preguntas y "
    "Análisis de Rendimiento con NumPy"
)
INSTITUCION = "UNIVERSIDAD TECNOLÓGICA DEL PERÚ"
FACULTAD = "Facultad de Ingeniería de Sistemas e Informática"
AUTORES = (
    "Luis Enrique Mamani Aguilar (U23259985)",
    "Eduardo Franco Cortez Benites (U1421099)",
    "Paul Williams Gallardo Villa (U1614474)",
)
AZUL_UTP = RGBColor(0x1E, 0x3A, 0x8A)
GRIS = RGBColor(0x33, 0x41, 0x55)
RAIZ = Path(__file__).resolve().parents[1]
FIGURAS = (
    ("Distribución de puntajes obtenidos en la simulación",
     RAIZ / "resultados" / "simulacion" / "puntajes_simulados.png"),
    ("Pantalla de bienvenida y captura del nombre",
     RAIZ / "docs" / "informe" / "figuras" / "01_bienvenida.png"),
    ("Presentación de una fase del torneo",
     RAIZ / "docs" / "informe" / "figuras" / "02_fase.png"),
    ("Pregunta con imagen y cuatro alternativas",
     RAIZ / "docs" / "informe" / "figuras" / "03_pregunta.png"),
    ("Retroalimentación visual después de responder",
     RAIZ / "docs" / "informe" / "figuras" / "04_retroalimentacion.png"),
    ("Pantalla final con clasificación de rendimiento",
     RAIZ / "docs" / "informe" / "figuras" / "05_resultado.png"),
)


def insertar_despues(parrafo, texto="", estilo=None):
    nuevo_elemento = OxmlElement("w:p")
    parrafo._p.addnext(nuevo_elemento)
    nuevo = Paragraph(nuevo_elemento, parrafo._parent)
    if texto:
        nuevo.add_run(texto)
    if estilo:
        nuevo.style = estilo
    return nuevo


def insertar_antes(parrafo, texto="", estilo=None):
    nuevo_elemento = OxmlElement("w:p")
    parrafo._p.addprevious(nuevo_elemento)
    nuevo = Paragraph(nuevo_elemento, parrafo._parent)
    if texto:
        nuevo.add_run(texto)
    if estilo:
        nuevo.style = estilo
    return nuevo


def eliminar_parrafo(parrafo):
    elemento = parrafo._element
    elemento.getparent().remove(elemento)
    parrafo._p = parrafo._element = None


def configurar_run(run, tamano, negrita=False, color=None):
    run.bold = negrita
    run.font.name = "Times New Roman"
    run.font.size = Pt(tamano)
    if color:
        run.font.color.rgb = color


def agregar_linea_ficha(cursor, etiqueta, texto):
    parrafo = insertar_despues(cursor)
    parrafo.paragraph_format.left_indent = Inches(0.8)
    parrafo.paragraph_format.first_line_indent = Inches(0)
    parrafo.paragraph_format.space_after = Pt(3)
    parrafo.paragraph_format.tab_stops.add_tab_stop(Inches(1.45))
    run_etiqueta = parrafo.add_run(etiqueta)
    configurar_run(run_etiqueta, 10, negrita=bool(etiqueta))
    run_texto = parrafo.add_run(f"\t{texto}")
    configurar_run(run_texto, 10)
    return parrafo


def agregar_numero_pagina(seccion):
    encabezado = seccion.header
    parrafo = encabezado.paragraphs[0]
    parrafo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = parrafo.add_run()
    inicio = OxmlElement("w:fldChar")
    inicio.set(qn("w:fldCharType"), "begin")
    instruccion = OxmlElement("w:instrText")
    instruccion.set(qn("xml:space"), "preserve")
    instruccion.text = " PAGE "
    final = OxmlElement("w:fldChar")
    final.set(qn("w:fldCharType"), "end")
    run._r.extend((inicio, instruccion, final))


def agregar_indice(parrafo):
    parrafo.paragraph_format.first_line_indent = Inches(0)
    inicio = OxmlElement("w:fldChar")
    inicio.set(qn("w:fldCharType"), "begin")
    inicio.set(qn("w:dirty"), "true")
    instruccion = OxmlElement("w:instrText")
    instruccion.set(qn("xml:space"), "preserve")
    instruccion.text = ' TOC \\o "1-3" \\h \\z \\u '
    separador = OxmlElement("w:fldChar")
    separador.set(qn("w:fldCharType"), "separate")
    texto = OxmlElement("w:t")
    texto.text = "Actualiza este campo para mostrar el índice."
    final = OxmlElement("w:fldChar")
    final.set(qn("w:fldCharType"), "end")
    run = parrafo.add_run()
    run._r.extend((inicio, instruccion, separador, texto, final))


def activar_actualizacion_de_campos(documento):
    configuracion = documento.settings.element
    actualizar = configuracion.find(qn("w:updateFields"))
    if actualizar is None:
        actualizar = OxmlElement("w:updateFields")
        configuracion.append(actualizar)
    actualizar.set(qn("w:val"), "true")


def agregar_portada_e_indice(documento):
    resumen = next(
        (parrafo for parrafo in documento.paragraphs
         if parrafo.text.strip() == "Resumen"), None)
    if resumen is None:
        raise ValueError("No se encontró el resumen para insertar la portada y el índice.")

    for parrafo in list(documento.paragraphs):
        if parrafo._p is resumen._p:
            break
        eliminar_parrafo(parrafo)

    institucion = insertar_antes(resumen, INSTITUCION)
    institucion.alignment = WD_ALIGN_PARAGRAPH.CENTER
    institucion.paragraph_format.space_before = Pt(45)
    institucion.paragraph_format.space_after = Pt(6)
    configurar_run(institucion.runs[0], 15, negrita=True)

    cursor = insertar_despues(institucion, FACULTAD)
    cursor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cursor.paragraph_format.space_after = Pt(80)
    configurar_run(cursor.runs[0], 11, color=GRIS)

    cursor = insertar_despues(cursor, TITULO)
    cursor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cursor.paragraph_format.space_after = Pt(18)
    configurar_run(cursor.runs[0], 21, negrita=True, color=AZUL_UTP)

    cursor = insertar_despues(cursor, "Informe técnico del proyecto")
    cursor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cursor.paragraph_format.space_after = Pt(5)
    configurar_run(cursor.runs[0], 12, color=GRIS)

    cursor = insertar_despues(cursor, "Proyecto final")
    cursor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cursor.paragraph_format.space_after = Pt(70)
    configurar_run(cursor.runs[0], 10.5, negrita=True)

    cursor = agregar_linea_ficha(
        cursor, "Integrantes:", "Mamani Aguilar, Luis Enrique  |  U23259985")
    cursor = agregar_linea_ficha(
        cursor, "", "Cortez Benites, Eduardo Franco  |  U1421099")
    cursor = agregar_linea_ficha(
        cursor, "", "Gallardo Villa, Paul Williams  |  U1614474")
    cursor = agregar_linea_ficha(
        cursor, "Asignatura:", "Programación de Computadoras (100000I42M)")
    cursor = agregar_linea_ficha(cursor, "Periodo:", "2026 - Ciclo 1 (marzo)")

    cursor = insertar_despues(cursor, "Lima, Perú")
    cursor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cursor.paragraph_format.space_before = Pt(72)
    cursor.paragraph_format.space_after = Pt(3)
    configurar_run(cursor.runs[0], 11)

    cursor = insertar_despues(cursor, "2026")
    cursor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    configurar_run(cursor.runs[0], 10.5, color=GRIS)

    salto = insertar_despues(cursor)
    salto.add_run().add_break(WD_BREAK.PAGE)

    cursor = insertar_despues(salto, "Índice")
    cursor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cursor.paragraph_format.space_after = Pt(24)
    configurar_run(cursor.runs[0], 18, negrita=True, color=AZUL_UTP)

    cursor = insertar_despues(cursor)
    agregar_indice(cursor)
    salto = insertar_despues(cursor)
    salto.add_run().add_break(WD_BREAK.PAGE)

    resumen.style = documento.styles["Heading 1"]
    activar_actualizacion_de_campos(documento)


def agregar_figuras(documento):
    parrafos = list(documento.paragraphs)
    for numero, (leyenda, ruta) in enumerate(FIGURAS, start=1):
        if not ruta.exists():
            raise FileNotFoundError(f"No se encontró la evidencia requerida: {ruta}")
        parrafo_leyenda = next(
            (parrafo for parrafo in parrafos if parrafo.text.strip() == leyenda), None)
        if parrafo_leyenda is None:
            raise ValueError(f"No se encontró la leyenda en el Word: {leyenda}")
        parrafo_imagen = insertar_antes(parrafo_leyenda)
        parrafo_imagen.alignment = WD_ALIGN_PARAGRAPH.CENTER
        parrafo_imagen.add_run().add_picture(str(ruta), width=Inches(6.4))
        parrafo_leyenda.text = f"Figura {numero}. {leyenda}"
        parrafo_leyenda.style = documento.styles["Caption"]


def agregar_encabezado_referencias(documento):
    referencia_inicial = next(
        (parrafo for parrafo in documento.paragraphs
         if "Array Programming with NumPy" in parrafo.text), None)
    if referencia_inicial is None:
        raise ValueError("No se encontró el inicio de las referencias bibliográficas.")
    encabezado = insertar_antes(referencia_inicial, "Referencias", "Heading 1")
    encabezado.alignment = WD_ALIGN_PARAGRAPH.CENTER


def completar_word(entrada, salida):
    documento = Document(entrada)
    seccion = documento.sections[0]
    seccion.page_width = Inches(8.27)
    seccion.page_height = Inches(11.69)
    seccion.top_margin = Inches(1)
    seccion.bottom_margin = Inches(1)
    seccion.left_margin = Inches(1)
    seccion.right_margin = Inches(1)
    seccion.different_first_page_header_footer = True
    agregar_numero_pagina(seccion)

    agregar_portada_e_indice(documento)
    agregar_figuras(documento)
    agregar_encabezado_referencias(documento)

    documento.core_properties.title = TITULO
    documento.core_properties.author = "; ".join(AUTORES)
    documento.core_properties.subject = "Proyecto final de Programación de Computadoras"
    Path(salida).parent.mkdir(parents=True, exist_ok=True)
    documento.save(salida)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--entrada", required=True)
    parser.add_argument("--salida", required=True)
    argumentos = parser.parse_args()
    completar_word(argumentos.entrada, argumentos.salida)
    print(f"Word APA 7 creado: {argumentos.salida}")


if __name__ == "__main__":
    main()
